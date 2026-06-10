"""Gerador de ATA de Defesa de TCC em formato .docx."""
from __future__ import annotations

import io
import os
import unicodedata
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers de formatação
# ---------------------------------------------------------------------------

_DIAS = [
    "", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito",
    "nove", "dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis",
    "dezessete", "dezoito", "dezenove", "vinte", "vinte e um", "vinte e dois",
    "vinte e três", "vinte e quatro", "vinte e cinco", "vinte e seis",
    "vinte e sete", "vinte e oito", "vinte e nove", "trinta", "trinta e um",
]

_MESES = [
    "", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
    "julho", "agosto", "setembro", "outubro", "novembro", "dezembro",
]

_CENTENAS = ["", "cem", "duzentos", "trezentos", "quatrocentos",
             "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
_DEZENAS  = ["", "dez", "vinte", "trinta", "quarenta", "cinquenta",
             "sessenta", "setenta", "oitenta", "noventa"]
_UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis",
             "sete", "oito", "nove"]
_TEENS    = ["dez", "onze", "doze", "treze", "quatorze", "quinze",
             "dezesseis", "dezessete", "dezoito", "dezenove"]


def _num_extenso(n: int) -> str:
    """Converte inteiro 1-9999 para extenso em português."""
    if n <= 0:
        return "zero"
    partes = []
    mil = n // 1000
    resto = n % 1000
    if mil:
        partes.append("mil" if mil == 1 else f"{_num_extenso(mil)} mil")
    cen = resto // 100
    dez = (resto % 100) // 10
    uni = resto % 10
    if cen:
        if dez == 0 and uni == 0:
            partes.append(_CENTENAS[cen])
        else:
            partes.append("cento" if cen == 1 else _CENTENAS[cen])
    if dez == 1:
        partes.append(_TEENS[uni])
    elif dez:
        partes.append(_DEZENAS[dez] + (f" e {_UNIDADES[uni]}" if uni else ""))
    elif uni:
        partes.append(_UNIDADES[uni])
    return " e ".join(partes) if partes else "zero"


def _hora_extenso(hora_str: str) -> tuple[str, str]:
    """Retorna (extenso, numerico) do horário 'HH:MM'."""
    try:
        h, _ = hora_str.split(":")
        h = int(h)
        return _num_extenso(h), str(h)
    except Exception:
        return hora_str, hora_str


def _data_extenso(data_str: str) -> tuple[str, str, str, str, str, str]:
    """Retorna (dia_ext, dia_num, mes_ext, ano_ext, ano_num, mes_num) de 'YYYY-MM-DD'."""
    try:
        ano, mes, dia = data_str.split("-")
        return (
            _num_extenso(int(dia)),
            str(int(dia)),
            _MESES[int(mes)],
            _num_extenso(int(ano)),
            ano,
            mes,
        )
    except Exception:
        return data_str, data_str, "", data_str, data_str, ""


def _titulo_abrev(titulo: str | None) -> str:
    """Converte 'Doutorado' → 'Dr.' etc."""
    mapa = {
        "Doutorado": "Dr.",
        "Mestrado": "M.Sc.",
        "Especialista": "Esp.",
        "Graduação": "",
    }
    return mapa.get(titulo or "", "")


def _prefixo(titulo: str | None) -> str:
    abrev = _titulo_abrev(titulo)
    return f"Prof. {abrev}".strip() if abrev else "Prof."


def _filiacao_label(tipo: str | None) -> str:
    return "FASI" if tipo == "Interno" else "Externo"


# ---------------------------------------------------------------------------
# Helpers de estilo para python-docx
# ---------------------------------------------------------------------------

def _set_paragraph_spacing(para, before: int = 0, after: int = 0, line: int = None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _add_run(para, text: str, bold: bool = False, size_pt: int = 12,
             font_name: str = "Times New Roman", color_rgb: tuple | None = None) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    # Força fonte para português (latin + cs)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rPr.insert(0, rFonts)


def _centered_para(doc: Document, text: str = "", bold: bool = False,
                   size_pt: int = 12, spacing_before: int = 0,
                   spacing_after: int = 0) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, before=spacing_before, after=spacing_after)
    if text:
        _add_run(para, text, bold=bold, size_pt=size_pt)
    return para


# ---------------------------------------------------------------------------
# Geração do documento
# ---------------------------------------------------------------------------

def gerar_ata_docx(requerimento: dict, funcionarios: list[dict]) -> bytes:
    """
    Gera a ATA de Defesa em .docx e retorna os bytes do arquivo.

    requerimento: dict com campos da tabela requerimento_tcc_submissions
    funcionarios: lista de dicts da tabela funcionarios (para buscar titulo/tipo)
    """
    doc = Document()

    # Margens: superior 2.5cm, inferior 2.5cm, esquerda 3cm, direita 2cm
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

    # -- Logo UFPA --
    logo_path = os.path.join(os.path.dirname(__file__), "..", "..", "..", "resources", "ufpa.png")
    logo_path = os.path.normpath(logo_path)
    if os.path.exists(logo_path):
        para_logo = doc.add_paragraph()
        para_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(para_logo, before=0, after=60)
        run_logo = para_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(2.5))

    # -- Cabeçalho institucional --
    for texto, bold, size in [
        ("SERVIÇO PÚBLICO FEDERAL", False, 12),
        ("UNIVERSIDADE FEDERAL DO PARÁ", False, 12),
        ("CAMPUS UNIVERSITÁRIO DO TOCANTINS - CAMETÁ", False, 12),
        ("FACULDADE DE SISTEMAS DE INFORMAÇÃO – FASI", True, 12),
    ]:
        _centered_para(doc, texto, bold=bold, size_pt=size, spacing_before=0, spacing_after=20)

    # -- Título do documento --
    _centered_para(doc, "", spacing_before=120, spacing_after=0)
    _centered_para(doc, "ATA DE DEFESA DE TRABALHO DE CONCLUSÃO DE CURSO",
                   bold=True, size_pt=12, spacing_before=0, spacing_after=200)

    # -- Dados da defesa --
    data_str   = requerimento.get("data_defesa") or ""
    hora_str   = requerimento.get("horario_defesa") or ""
    local_str  = requerimento.get("local_defesa") or "Google Meet"
    nome_aluno = (requerimento.get("nome_aluno") or "").upper()
    matricula  = requerimento.get("matricula") or ""
    turma      = requerimento.get("turma") or ""
    titulo_tcc = requerimento.get("titulo_trabalho") or ""

    dia_ext, dia_num, mes_ext, ano_ext, ano_num, _ = _data_extenso(data_str)
    hora_ext, hora_num = _hora_extenso(hora_str)

    # -- Busca título/tipo de cada membro na tabela funcionários --
    func_map: dict[str, dict] = {
        f["nome"].strip().lower(): f for f in funcionarios
    }

    def _get_func(nome: str | None) -> dict:
        if not nome:
            return {}
        return func_map.get(nome.strip().lower(), {})

    orientador_nome = requerimento.get("orientador") or ""
    m1_nome = requerimento.get("membro_banca1") or ""
    m2_nome = requerimento.get("membro_banca2") or ""
    m3_nome = requerimento.get("membro_banca3") or ""

    ori_f  = _get_func(orientador_nome)
    m1_f   = _get_func(m1_nome)
    m2_f   = _get_func(m2_nome)
    m3_f   = _get_func(m3_nome) if m3_nome and m3_nome.lower() not in ("", "nenhum") else {}

    ori_prefix = _prefixo(ori_f.get("titulo"))
    m1_prefix  = _prefixo(m1_f.get("titulo"))
    m2_prefix  = _prefixo(m2_f.get("titulo"))
    m3_prefix  = _prefixo(m3_f.get("titulo")) if m3_f else ""

    # -- Corpo da ATA (parágrafo justificado) --
    para_body = doc.add_paragraph()
    para_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para_body, before=0, after=240, line=276)

    # Constrói a lista de membros para o texto
    membros_texto = f"{m1_prefix} {m1_nome}".strip() + f" e {m2_prefix} {m2_nome}".strip()
    if m3_f and m3_nome:
        membros_texto += f", e {m3_prefix} {m3_nome}".strip()

    _add_run(para_body,
        f"Aos {dia_ext} ({dia_num}) dias do mês de {mes_ext} de {ano_ext} ({ano_num}), "
        f"às {hora_ext} ({hora_num}h) reuniu-se via {local_str}, a Banca Examinadora para "
        f"Avaliação do Trabalho de Conclusão de Curso – TCC, composta pelo "
        f"{ori_prefix} {orientador_nome}, orientador e os membros {membros_texto}, "
        f"para avaliarem o trabalho Intitulado: "
    )
    _add_run(para_body, titulo_tcc, bold=True)
    _add_run(para_body, f", do discente ")
    _add_run(para_body, f"{nome_aluno} ({matricula})", bold=True)
    _add_run(para_body,
        f", do Curso Sistemas de Informação - Turma {turma}. "
        f"Após a apresentação e defesa, a banca em ato de compilação final, "
        f"considerando a análise das fichas de avaliação e de acordo como Regimento "
        f"Geral da UFPA no Capítulo IV, art. 178, considerou o trabalho "
    )
    _add_run(para_body, "[APROVADO|REPROVADO]", bold=True, color_rgb=(255, 0, 0))
    _add_run(para_body, " e atribuiu o conceito ")
    _add_run(para_body, "[INSUFICIENTE|REGULAR|BOM|EXCELENTE]", bold=True, color_rgb=(255, 0, 0))
    _add_run(para_body,
        ", considerando as sugestões necessárias ao processo de homologação. "
        "Nada mais havendo a tratar, foi lavrada a presente ATA que, depois de lida "
        "e aprovada, vai assinada pelos professores que compõem a Banca Examinadora "
        "e pelo aluno."
    )

    # -- Bloco de assinaturas --
    def _sig(nome: str, prefix: str, papel: str, filiacao: str):
        if not nome or nome.lower() in ("", "nenhum"):
            return
        doc.add_paragraph()  # espaço
        p_nome = _centered_para(doc, f"{prefix} {nome}".strip(), bold=False, size_pt=12,
                                 spacing_before=180, spacing_after=0)
        _centered_para(doc, papel + (" – " + filiacao if filiacao else ""),
                       bold=False, size_pt=12, spacing_before=0, spacing_after=0)

    _sig(orientador_nome, ori_prefix, "Orientador", _filiacao_label(ori_f.get("tipo")))
    _sig(m1_nome, m1_prefix, "Membro", _filiacao_label(m1_f.get("tipo")))
    _sig(m2_nome, m2_prefix, "Membro", _filiacao_label(m2_f.get("tipo")))
    if m3_f and m3_nome:
        _sig(m3_nome, m3_prefix, "Membro", _filiacao_label(m3_f.get("tipo")))

    # Aluno
    doc.add_paragraph()
    _centered_para(doc, (requerimento.get("nome_aluno") or "").title(),
                   bold=False, size_pt=12, spacing_before=180, spacing_after=0)
    _centered_para(doc, "Aluno", bold=False, size_pt=12, spacing_before=0, spacing_after=0)

    # Serializa para bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Envio de e-mail com a ATA
# ---------------------------------------------------------------------------

_SIPAC_INSTRUCOES = """\
Prezado(a) orientador(a),

Segue em anexo a Ata de Defesa de TCC gerada automaticamente pelo sistema FasiTech.

Para cadastrar a ATA no SIPAC, siga os passos abaixo:

Passo 1. Entre no SIPAC
Passo 2. Cadastre Documento
Passo 3. Na tela de cadastro:
  Passo 3.1 Tipo de Documento: ATA DE DEFESA DE TRABALHO DE CONCLUSÃO DE CURSO (TCC)
  Passo 3.2 Natureza: Ostensivo
  Passo 3.3 Assunto Detalhado: ATA DE DEFESA DE TRABALHO DE CONCLUSÃO DE CURSO DO ALUNO {nome_aluno}
  Passo 3.4 Forma de Documento: Anexar documento
  Passo 3.5 Ao preencher dados referente ao documento anexado, marque os envolvidos.
    Passo 3.5.1 Caso todos os envolvidos sejam da UFPA, procure e marque cada um.
    Passo 3.5.2 Caso algum dos envolvidos seja externo da UFPA, você pode solicitar \
assinatura sou.gov ou cadastrar via SIGAA.
  Passo 3.6 Aperte Continuar e posteriormente em Continuar
  Passo 3.7 Na seção "Informar Interessados no Documento": coloque orientador e aluno como interessados.
  Passo 3.8 Na seção "Informar Dados da Movimentação Inicial", como Unidade de Destino \
sempre coloque 11.19.36, que corresponde à FACULDADE DE SISTEMAS DE INFORMAÇÃO - CAMETÁ (11.19.36).
  Passo 3.9 Faça conferência das informações e finalize o processo.

Atenciosamente,
FasiTech – Sistema de Gestão Acadêmica FASI
"""


def enviar_ata_por_email(requerimento: dict, funcionarios: list[dict]) -> None:
    """Gera a ATA e envia por e-mail ao orientador com instruções SIPAC."""
    from backend.infrastructure.database.repository import get_funcionario_emails_by_nomes
    from backend.infrastructure.email.service import send_email_with_attachments

    orientador_nome = requerimento.get("orientador") or ""
    nome_aluno = requerimento.get("nome_aluno") or ""

    emails = get_funcionario_emails_by_nomes([orientador_nome]) if orientador_nome else []
    if not emails:
        raise ValueError(f"E-mail do orientador '{orientador_nome}' não encontrado na base de dados.")

    docx_bytes = gerar_ata_docx(requerimento, funcionarios)
    # Normaliza acentos para nome seguro de arquivo
    nome_sem_acento = unicodedata.normalize('NFKD', nome_aluno.upper())
    nome_sem_acento = ''.join(c for c in nome_sem_acento if not unicodedata.combining(c))
    nome_arquivo = nome_sem_acento.replace(' ', '_')
    filename = f"ATA_DEFESA_{nome_arquivo}.doc"

    corpo = _SIPAC_INSTRUCOES.format(nome_aluno=nome_aluno.upper())
    assunto = f"ATA de Defesa de TCC – {nome_aluno}"

    # send_email_with_attachments aceita attachments como lista de (filename, bytes, mimetype)
    send_email_with_attachments(
        subject=assunto,
        body=corpo,
        recipients=emails,
        attachments=[(filename, docx_bytes, "application/msword")],
    )
